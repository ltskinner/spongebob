
import os
import sys
import random

import docx2txt
from docx import Document
from docx.shared import Inches

from cv2 import cv2

import argparse

import os.path
banner_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 
                           'banner.txt')
f = open(banner_path, 'r', encoding='utf-8')
print(f.read())
f.close()


def spongebobify(text=None,
                 file_to_convert=None,
                 save_picture=True,
                 save_docx=False) -> str:
    if text is None and file_to_convert is None:
        raw_text = input("EntEr The TexT YOu WaNT tO cONvERt:\n")
        file_to_convert = "spongebob"

    elif file_to_convert is not None:
        if '.docx' not in file_to_convert and '.txt' not in file_to_convert:
            print("sPEcIFy .dOcX or .Txt AnD rUn aGaiN")
            sys.exit(0)

        if file_to_convert in os.listdir():
            if '.docx' in file_to_convert:
                raw_text = docx2txt.process(file_to_convert)
            elif '.txt' in file_to_convert:
                with open(file_to_convert, 'r') as file:
                    raw_text = file.read()
        else:
            print("FIle DoeS NOt EXisT")
            sys.exit(0)

    elif text is not None:
        raw_text = text
        file_to_convert = "spongebob"

    sponge_text = to_spongebob_text(raw_text)
    partitioned_text = split_text(sponge_text)

    if save_picture:
        create_and_write_meme(partitioned_text)

    if save_docx:
        sponge_file = name_docx_file(file_to_convert)
        write_to_docx(partitioned_text, sponge_file)

    print(sponge_text)

    return sponge_text


def to_spongebob_text(raw_text):
    raw_text = raw_text.lower()
    length = 2 + int(len(raw_text)/2)
    cases = [0]*length + [1]*length
    random.shuffle(cases)

    last, prev = 0, 0
    for i in range(len(cases)):
        if cases[i] == last and cases[i] == prev:  # Three in a row
            cases[i] = 1 - cases[i]

        last = prev
        prev = cases[i]

    sponge_text = ''
    for i in range(len(raw_text)):
        if cases[i] == 1:
            sponge_text += raw_text[i].upper()
        else:
            sponge_text += raw_text[i]

    return sponge_text


def create_and_write_meme(partitioned_text):
    my_path = os.path.abspath(os.path.dirname(__file__))
    pic_path = os.path.join(my_path, 'spongebob.jpg')

    # height, width, channels
    img = cv2.imread(pic_path)
    img_height, img_width, _ = img.shape

    print('[!] Use num characters to determine font')
    # def find_font_size()

    top_text, bot_text = partitioned_text.split("|SPONGEPIC|")

    # TODO: swap font type -> IMPACT, use Dank Learnings meme writer maybe?
    font = find_font_params(top_text, bot_text)
    label_color = (255, 255, 255) #255 #(0, 0, 0)
    # TODO: add black outline to text

    (top_label_width, top_label_height), _ = cv2.getTextSize(top_text,
                                                             font['font'],
                                                             font['scale'],
                                                             font['thickness'])
    top_center = top_label_width / 2
    top_left = int((img_width/2) - top_center)
    top_top = int((img_height * .05) + top_label_height)  # 5% down

    (bot_label_width, bot_label_height), _ = cv2.getTextSize(bot_text,
                                                             font['font'],
                                                             font['scale'],
                                                             font['thickness'])
    bot_center = bot_label_width / 2
    bot_left = int((img_width/2) - bot_center)
    bot_bot = int((img_height * .95) - bot_label_height)  # 5% down

    label_color = (255, 255, 255) #255 #(0, 0, 0)
    cv2.putText(img, top_text, (top_left, top_top), font['font'], font['scale'], label_color, thickness=font['thickness']S)
    cv2.putText(img, bot_text, (bot_left, bot_bot), font['font'], font['scale'], label_color, thickness=font['thickness'])

    cv2.imwrite(os.path.join(os.getcwd(), 'meme.jpg'), img)


def find_font_params(top_text, bot_text):
    FONT = cv2.FONT_HERSHEY_COMPLEX
    FONT_SCALE = 1.5
    FONT_THICKNESS = 2

    (top_label_width, top_label_height), _ = cv2.getTextSize(top_text,
                                                             FONT,
                                                             FONT_SCALE,
                                                             FONT_THICKNESS)
    (bot_label_width, bot_label_height), _ = cv2.getTextSize(bot_text,
                                                             FONT,
                                                             FONT_SCALE,
                                                             FONT_THICKNESS)

    max_width = max([top_label_width, bot_label_width])

    font = {
        'font': FONT,
        'scale': FONT_SCALE,
        'thickness': FONT_THICKNESS
    }
    return font

def name_docx_file(file_to_convert):
    file_to_convert = file_to_convert.replace('.', "_spongebob.")
    return str(to_spongebob_text(file_to_convert.split('.')[0]) + ".docx")


def write_to_docx(docx_text, sponge_file):
    my_path = os.path.abspath(os.path.dirname(__file__))
    pic_path = os.path.join(my_path, 'spongebob.jpg')
    text_segs = docx_text.split("|SPONGEPIC|")

    document = Document()
    document.add_paragraph(text_segs[0])
    document.add_picture(pic_path, width=Inches(6))
    document.add_paragraph(text_segs[1])

    document.save(sponge_file)


def split_text(sponge_text):
    newlines = sponge_text.split('\n')
    if len(newlines) > 2:
        newlines.insert(int(len(newlines)/2), "|SPONGEPIC|")
    elif len(newlines) == 1:
        newwords = newlines[0].split(' ')
        if len(newwords) > 2:
            newwords.insert(int(len(newwords)/2), "|SPONGEPIC|")
        else:
            newwords.append("|SPONGEPIC|")
        newlines[0] = ' '.join(newwords)

    return '\n'.join(newlines)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Convert text into SpOngEBOb TExT")
    parser.add_argument("--file_to_convert", help="TeXt fiLe to CoNVerT")
    args = parser.parse_args()

    spongebobify(file_to_convert=args.file_to_convert)
